#!/usr/bin/env python3
"""Generate writing-under-surveillance.docx by parsing index.html."""

import re
import html as html_mod
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Parse HTML ────────────────────────────────────────────────────────────────
with open('/home/milwrite/Quimbot/docs/writing-under-surveillance/index.html', encoding='utf-8') as f:
    src = f.read()

def strip_tags(text):
    """Remove HTML tags and decode entities."""
    text = re.sub(r'<cite>[^<]*</cite>', lambda m: re.sub(r'<[^>]+>', '', m.group()), text)
    text = re.sub(r'<em>([^<]*)</em>', r'\1', text)
    text = re.sub(r'<a[^>]*>([^<]*)</a>', r'\1', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = html_mod.unescape(text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def strip_tags_keep_em(text):
    """Return list of (text, is_italic) tuples, preserving <em> spans."""
    def clean_segment(s):
        """Strip tags and entities but preserve inter-word spaces."""
        s = re.sub(r'<cite>([^<]*)</cite>', lambda m: re.sub(r'<[^>]+>', '', m.group()), s)
        s = re.sub(r'<em>([^<]*)</em>', r'\1', s)
        s = re.sub(r'<a[^>]*>([^<]*)</a>', r'\1', s)
        s = re.sub(r'<[^>]+>', '', s)
        s = html_mod.unescape(s)
        s = re.sub(r'\s+', ' ', s)
        return s  # do NOT strip — preserve leading/trailing spaces between runs

    result = []
    pos = 0
    for m in re.finditer(r'<em>(.*?)</em>', text, re.DOTALL):
        before = clean_segment(text[pos:m.start()])
        if before.strip():
            result.append((before, False))
        em_text = clean_segment(m.group(1)).strip()
        if em_text:
            result.append((em_text, True))
            # ensure space after italic run if the next char is not punctuation
            after_start = text[m.end():m.end()+1]
            if after_start and after_start not in (',', '.', ';', ':', ')', '—', '\n'):
                result.append((' ', False))
        pos = m.end()
    tail = clean_segment(text[pos:]).strip()
    if tail:
        result.append((tail, False))
    return result if result else [(strip_tags(text), False)]

# Extract article paragraphs (between <article> tags)
article_match = re.search(r'<article>(.*?)</article>', src, re.DOTALL)
article_html = article_match.group(1) if article_match else ''

# Extract epigraph
epigraph_match = re.search(r'<div class="pullquote">(.*?)</div>', article_html, re.DOTALL)
epigraph_text = strip_tags(epigraph_match.group(1)) if epigraph_match else ''

# Extract body paragraphs (stop before .revision-tag and addendum div)
body_section = re.split(r'<div class="revision-tag">', article_html)[0]
body_paras_raw = re.findall(r'<p>(.*?)</p>', body_section, re.DOTALL)

# Extract works cited items
wc_match = re.search(r'<div class="works-cited">(.*?)</div>\s*</div>', src, re.DOTALL)
wc_items_raw = re.findall(r'<li>(.*?)</li>', wc_match.group(1) if wc_match else '', re.DOTALL)

# ── Build Document ────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.25)
    section.bottom_margin = Inches(1.25)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

style = doc.styles['Normal']
style.font.name = 'Georgia'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(10)
style.paragraph_format.line_spacing = Pt(20)

def body_para(parts):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    for text, italic in parts:
        r = p.add_run(text)
        r.italic = italic
    return p

# Title
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Writing Under Surveillance')
r.bold = True
r.font.size = Pt(18)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run('The Problem with AI Detection')
r.font.size = Pt(14)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(20)
r = p.add_run('Zach Muhlbauer  \u00b7  Draft v3 \u00b7 Feb 2026')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(100, 100, 100)

# Epigraph
if epigraph_text:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(18)
    r = p.add_run(epigraph_text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(80, 80, 80)

# Body paragraphs
for raw in body_paras_raw:
    # skip empty, byline, or addendum paragraphs
    clean = strip_tags(raw).strip()
    if not clean:
        continue
    if 'font-size' in raw or 'companion resource' in raw:
        continue
    parts = strip_tags_keep_em(raw)
    body_para(parts)

# Works Cited
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run('Works Cited')
r.bold = True
r.font.size = Pt(13)

for raw in wc_items_raw:
    clean = strip_tags(raw).strip()
    if not clean:
        continue
    parts = strip_tags_keep_em(raw)
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.25)
    for text, italic in parts:
        r = p.add_run(text)
        r.italic = italic

# Save
out = '/home/milwrite/Quimbot/docs/writing-under-surveillance/writing-under-surveillance.docx'
doc.save(out)
print(f'Saved: {out}')
