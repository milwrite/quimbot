#!/usr/bin/env python3
"""Convert Writing Under Surveillance HTML pages to .docx files."""

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re, os

BASE = "/home/milwrite/Quimbot/docs/writing-under-surveillance"
OUT  = "/home/milwrite/Quimbot/docs/writing-under-surveillance"

# ── helpers ─────────────────────────────────────────────────────────────────

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Georgia"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text, style="Normal", size=12, bold=False,
                  italic=False, space_before=0, space_after=8,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = alignment
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def inline_text(el):
    """Flatten an element to plain text, stripping HTML tags."""
    return el.get_text(" ", strip=True)

def add_rich_paragraph(doc, el, size=12, space_before=0, space_after=8):
    """Add a paragraph that preserves <em> and <cite> italics."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for child in el.children:
        if isinstance(child, NavigableString):
            t = str(child)
            if t.strip() or t == " ":
                run = p.add_run(t)
                set_font(run, size=size)
        elif child.name in ("em", "cite", "i"):
            run = p.add_run(child.get_text())
            set_font(run, size=size, italic=True)
        elif child.name == "strong":
            run = p.add_run(child.get_text())
            set_font(run, size=size, bold=True)
        elif child.name == "a":
            run = p.add_run(child.get_text())
            set_font(run, size=size)
        else:
            run = p.add_run(child.get_text())
            set_font(run, size=size)
    return p

# ── Article ──────────────────────────────────────────────────────────────────

def convert_article():
    with open(f"{BASE}/index.html", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Title
    h1 = soup.find("h1")
    title_text = h1.get_text(strip=True) if h1 else "Writing Under Surveillance"
    p = add_paragraph(doc, title_text, size=18, bold=True,
                      space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Byline
    byline = soup.find("p", class_="byline")
    if byline:
        author = byline.find("span", class_="author")
        badge  = byline.find("span", class_="badge")
        line   = ""
        if author: line += author.get_text(strip=True)
        if badge:  line += f"   {badge.get_text(strip=True)}"
        add_paragraph(doc, line, size=10, italic=True,
                      space_before=0, space_after=16)

    # Epigraph / pullquote
    pullquote = soup.find("div", class_="pullquote")
    if pullquote:
        # Quote text (everything before the <cite>)
        cite_el = pullquote.find("cite")
        cite_text = cite_el.get_text(strip=True) if cite_el else ""
        # Get full text then strip the cite portion
        full = pullquote.get_text(" ", strip=True)
        quote_text = full.replace(cite_text, "").strip().strip("—").strip()

        p = doc.add_paragraph()
        p.paragraph_format.left_indent   = Inches(0.5)
        p.paragraph_format.right_indent  = Inches(0.5)
        p.paragraph_format.space_before  = Pt(8)
        p.paragraph_format.space_after   = Pt(4)
        run = p.add_run(quote_text)
        set_font(run, size=11, italic=True)

        if cite_text:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent  = Inches(0.5)
            p2.paragraph_format.right_indent = Inches(0.5)
            p2.paragraph_format.space_after  = Pt(16)
            run = p2.add_run(f"\u2014 {cite_text}")
            set_font(run, size=10, italic=True)

    # Article body paragraphs
    article = soup.find("article")
    if article:
        for el in article.children:
            if not isinstance(el, Tag):
                continue
            if el.name == "p":
                add_rich_paragraph(doc, el, size=12, space_after=10)
            elif el.name == "div" and "pullquote" in el.get("class", []):
                pass  # already handled
            elif el.name == "div" and "revision-tag" in el.get("class", []):
                add_paragraph(doc, el.get_text(strip=True), size=9,
                              italic=True, space_before=12, space_after=4)
            elif el.name == "div":
                # addendum
                p_els = el.find_all("p")
                for p_el in p_els:
                    add_rich_paragraph(doc, p_el, size=11,
                                       space_before=14, space_after=4)

    # Works Cited
    wc_div = soup.find("div", class_="works-cited")
    if wc_div:
        doc.add_paragraph()
        add_paragraph(doc, "Works Cited", size=11, bold=True,
                      space_before=16, space_after=8)
        items = wc_div.find_all("li")
        for item in items:
            text = item.get_text(" ", strip=True)
            # Clean extra whitespace
            text = re.sub(r"\s+", " ", text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent       = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after        = Pt(4)
            run = p.add_run(text)
            set_font(run, size=10)

    out_path = f"{OUT}/writing-under-surveillance.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


# ── Teaching Tips ────────────────────────────────────────────────────────────

def convert_tips():
    with open(f"{BASE}/teaching-tips.html", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "html.parser")

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Title
    h1 = soup.find("h1")
    add_paragraph(doc, h1.get_text(strip=True) if h1 else "Small Wins & Teaching Tips",
                  size=18, bold=True, space_after=4)

    # Subtitle / byline
    byline = soup.find("p", class_="byline")
    if byline:
        add_paragraph(doc, byline.get_text(" ", strip=True), size=10,
                      italic=True, space_after=16)

    # Opening
    opening = soup.find("p", class_="opening")
    if opening:
        add_rich_paragraph(doc, opening, size=12, space_before=0, space_after=14)

    # Tips
    tips = soup.find_all("div", class_="tip")
    for tip in tips:
        title_el = tip.find("div", class_="tip-title")
        if title_el:
            add_paragraph(doc, title_el.get_text(strip=True), size=12,
                          bold=True, space_before=12, space_after=4)

        for child in tip.children:
            if not isinstance(child, Tag):
                continue
            if child.name == "p":
                add_rich_paragraph(doc, child, size=12, space_after=6)
            elif child.name == "div" and "example" in child.get("class", []):
                # Label
                label = child.find("strong")
                label_text = label.get_text(strip=True) if label else "Example"
                # Body text: everything after the <strong>
                body_parts = []
                for node in child.children:
                    if isinstance(node, NavigableString):
                        t = str(node).strip()
                        if t:
                            body_parts.append(t)
                    elif node.name != "strong":
                        body_parts.append(node.get_text(" ", strip=True))
                body = " ".join(body_parts).strip()

                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Inches(0.35)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(4)
                r1 = p.add_run(f"{label_text}: ")
                set_font(r1, size=10, bold=True, italic=True)
                r2 = p.add_run(body)
                set_font(r2, size=10, italic=True)

    # Footer note
    footer = soup.find("div", class_="footer-links")
    if footer:
        links = [a.get_text(strip=True) for a in footer.find_all("a")]
        add_paragraph(doc, " · ".join(links), size=9, italic=True,
                      space_before=20, space_after=0)

    out_path = f"{OUT}/small-wins-teaching-tips.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    convert_article()
    convert_tips()
    print("Done.")
